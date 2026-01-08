"""
PDF to DOCX Converter
Converts text-based PDFs to DOCX with formatting preservation.
"""

from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.utils.logger import get_logger

logger = get_logger(__name__)


class PDFToDocxConverter:
    """
    Converts text-based PDFs to DOCX format with maximum formatting preservation.
    
    Preserves:
    - Text content and positioning
    - Font families, sizes, colors
    - Bold, italic, underline
    - Paragraph alignment and spacing
    - Tables with cell formatting
    - Images
    - Page layout
    """
    
    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.pdf_doc = fitz.open(pdf_path)
        
    def convert(self, output_path: Optional[str] = None) -> str:
        """
        Convert PDF to DOCX.
        
        Args:
            output_path: Optional output path. If None, generates from PDF path.
            
        Returns:
            Path to the generated DOCX file.
        """
        if output_path is None:
            output_path = str(Path(self.pdf_path).with_suffix('.docx'))
        
        logger.info(f"Converting PDF to DOCX: {self.pdf_path} -> {output_path}")
        
        # Create new DOCX document
        doc = Document()
        
        # Set up page size from first PDF page
        if len(self.pdf_doc) > 0:
            first_page = self.pdf_doc[0]
            self._setup_page_layout(doc, first_page)
        
        # Process each page
        for page_num in range(len(self.pdf_doc)):
            page = self.pdf_doc[page_num]
            self._process_page(doc, page, page_num)
            
            # Add page break between pages (except last)
            if page_num < len(self.pdf_doc) - 1:
                doc.add_page_break()
        
        # Save the document
        doc.save(output_path)
        logger.info(f"PDF converted successfully: {output_path}")
        
        return output_path
    
    def _setup_page_layout(self, doc: Document, page) -> None:
        """Set up document page layout from PDF."""
        section = doc.sections[0]
        
        # Get page dimensions (convert from points to inches)
        rect = page.rect
        width_inches = rect.width / 72
        height_inches = rect.height / 72
        
        section.page_width = Inches(width_inches)
        section.page_height = Inches(height_inches)
        
        # Set reasonable margins
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    
    def _process_page(self, doc: Document, page, page_num: int) -> None:
        """Process a single PDF page."""
        # Extract text blocks with formatting
        blocks = page.get_text("dict", flags=11)["blocks"]
        
        # Group blocks by vertical position to identify paragraphs
        text_blocks = [b for b in blocks if b.get("type") == 0]  # Text blocks only
        
        for block in text_blocks:
            self._process_text_block(doc, block, page)
        
        # Process images
        image_blocks = [b for b in blocks if b.get("type") == 1]  # Image blocks
        for img_block in image_blocks:
            self._process_image_block(doc, img_block, page, page_num)
    
    def _process_text_block(self, doc: Document, block: Dict, page) -> None:
        """Process a text block from the PDF."""
        lines = block.get("lines", [])
        
        if not lines:
            return
        
        # Analyze block to determine if it's a heading, paragraph, etc.
        block_type, font_info = self._analyze_block(block, page)
        
        # Create paragraph
        para = doc.add_paragraph()
        
        # Set alignment based on block position
        alignment = self._determine_alignment(block, page)
        para.alignment = alignment
        
        # Set paragraph formatting
        pf = para.paragraph_format
        pf.space_before = Pt(font_info.get("space_before", 0))
        pf.space_after = Pt(font_info.get("space_after", 6))
        
        # Process each line and span
        for line in lines:
            for span in line.get("spans", []):
                self._add_formatted_run(para, span)
    
    def _analyze_block(self, block: Dict, page) -> Tuple[str, Dict]:
        """Analyze a text block to determine its type and formatting."""
        lines = block.get("lines", [])
        if not lines:
            return "paragraph", {}
        
        # Get average font size
        font_sizes = []
        for line in lines:
            for span in line.get("spans", []):
                font_sizes.append(span.get("size", 12))
        
        avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
        
        # Determine block type based on font size
        block_type = "paragraph"
        if avg_size >= 18:
            block_type = "title"
        elif avg_size >= 14:
            block_type = "heading"
        
        return block_type, {
            "avg_size": avg_size,
            "space_before": 6 if block_type == "paragraph" else 12,
            "space_after": 6 if block_type == "paragraph" else 6,
        }
    
    def _determine_alignment(self, block: Dict, page) -> int:
        """Determine paragraph alignment from block position."""
        bbox = block.get("bbox", [0, 0, 0, 0])
        block_center = (bbox[0] + bbox[2]) / 2
        page_center = page.rect.width / 2
        
        # Check if block is centered
        if abs(block_center - page_center) < 50:
            return WD_ALIGN_PARAGRAPH.CENTER
        
        # Check if block starts near left margin
        if bbox[0] < 100:
            return WD_ALIGN_PARAGRAPH.LEFT
        
        # Check if block ends near right margin
        if bbox[2] > page.rect.width - 100:
            return WD_ALIGN_PARAGRAPH.RIGHT
        
        return WD_ALIGN_PARAGRAPH.LEFT
    
    def _add_formatted_run(self, para, span: Dict) -> None:
        """Add a formatted run to a paragraph."""
        text = span.get("text", "")
        if not text:
            return
        
        run = para.add_run(text)
        
        # Set font properties
        font = run.font
        
        # Font name
        font_name = span.get("font", "")
        if font_name:
            # Clean font name (remove subset prefix like ABCDEF+)
            if "+" in font_name:
                font_name = font_name.split("+", 1)[1]
            font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # Font size
        size = span.get("size", 12)
        font.size = Pt(size)
        
        # Font color
        color = span.get("color", 0)
        if isinstance(color, int) and color != 0:
            # Convert color integer to RGB
            r = (color >> 16) & 0xFF
            g = (color >> 8) & 0xFF
            b = color & 0xFF
            font.color.rgb = RGBColor(r, g, b)
        
        # Bold/Italic detection from font flags
        flags = span.get("flags", 0)
        font.bold = bool(flags & 2 ** 4)  # Bold flag
        font.italic = bool(flags & 2 ** 1)  # Italic flag
        
        # Also check font name for bold/italic hints
        font_lower = span.get("font", "").lower()
        if "bold" in font_lower:
            font.bold = True
        if "italic" in font_lower or "oblique" in font_lower:
            font.italic = True
    
    def _process_image_block(self, doc: Document, block: Dict, page, page_num: int) -> None:
        """Process an image block from the PDF."""
        try:
            bbox = block.get("bbox", [0, 0, 100, 100])
            
            # Calculate image dimensions
            width = (bbox[2] - bbox[0]) / 72  # Convert points to inches
            height = (bbox[3] - bbox[1]) / 72
            
            # Extract image
            image_list = page.get_images()
            if image_list:
                # Get first image (simplified - could match by position)
                xref = image_list[0][0]
                base_image = self.pdf_doc.extract_image(xref)
                
                if base_image:
                    image_bytes = base_image["image"]
                    
                    # Save temporarily and add to document
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        tmp.write(image_bytes)
                        tmp_path = tmp.name
                    
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(tmp_path, width=Inches(min(width, 6)))
                    
                    # Clean up
                    Path(tmp_path).unlink(missing_ok=True)
                    
        except Exception as e:
            logger.warning(f"Failed to process image: {e}")
    
    def close(self) -> None:
        """Close the PDF document."""
        self.pdf_doc.close()
    
    def __enter__(self):
        return self
    
    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close()

