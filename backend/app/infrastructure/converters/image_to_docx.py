"""
Image to DOCX Converter
Converts images and scanned PDFs to DOCX using OCR.
"""

from typing import Optional, List, Dict, Any, Tuple
from pathlib import Path
import tempfile
import fitz  # PyMuPDF
from PIL import Image
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.infrastructure.ocr import OCREngine, ImagePreprocessor
from app.utils.logger import get_logger

logger = get_logger(__name__)


class ImageToDocxConverter:
    """
    Converts images and scanned PDFs to DOCX using OCR.
    
    Pipeline:
    1. Preprocess image (grayscale, threshold, deskew)
    2. Run OCR with bounding boxes
    3. Analyze layout (headings, paragraphs, etc.)
    4. Reconstruct document in DOCX format
    """
    
    def __init__(self):
        self.ocr_engine = OCREngine()
        self.preprocessor = ImagePreprocessor()
    
    def convert_image(self, image_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert a single image to DOCX.
        
        Args:
            image_path: Path to the image file.
            output_path: Optional output path.
            
        Returns:
            Path to the generated DOCX file.
        """
        if output_path is None:
            output_path = str(Path(image_path).with_suffix('.docx'))
        
        logger.info(f"Converting image to DOCX: {image_path}")
        
        # Load and preprocess image
        image = Image.open(image_path)
        preprocessed = self.preprocessor.preprocess_for_ocr(np.array(image))
        
        # Run OCR
        ocr_result = self.ocr_engine.extract_with_layout(preprocessed)
        
        # Create DOCX
        doc = Document()
        self._build_document_from_ocr(doc, ocr_result, image.size)
        
        doc.save(output_path)
        logger.info(f"Image converted successfully: {output_path}")
        
        return output_path
    
    def convert_scanned_pdf(self, pdf_path: str, output_path: Optional[str] = None) -> str:
        """
        Convert a scanned PDF to DOCX.
        
        Args:
            pdf_path: Path to the scanned PDF file.
            output_path: Optional output path.
            
        Returns:
            Path to the generated DOCX file.
        """
        if output_path is None:
            output_path = str(Path(pdf_path).with_suffix('.docx'))
        
        logger.info(f"Converting scanned PDF to DOCX: {pdf_path}")
        
        doc = Document()
        pdf_doc = fitz.open(pdf_path)
        
        try:
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                
                # Convert page to image
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for better OCR
                image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Preprocess
                preprocessed = self.preprocessor.preprocess_for_ocr(np.array(image))
                
                # OCR
                ocr_result = self.ocr_engine.extract_with_layout(preprocessed)
                
                # Build document content
                self._build_document_from_ocr(doc, ocr_result, image.size)
                
                # Add page break (except for last page)
                if page_num < len(pdf_doc) - 1:
                    doc.add_page_break()
        
        finally:
            pdf_doc.close()
        
        doc.save(output_path)
        logger.info(f"Scanned PDF converted successfully: {output_path}")
        
        return output_path
    
    def _build_document_from_ocr(self, doc: Document, ocr_result: Dict, image_size: Tuple[int, int]) -> None:
        """Build DOCX content from OCR results."""
        blocks = ocr_result.get("blocks", [])
        
        if not blocks:
            logger.warning("No text blocks found in OCR result")
            return
        
        # Sort blocks by vertical position (top to bottom)
        blocks = sorted(blocks, key=lambda b: b.get("bbox", [0, 0, 0, 0])[1])
        
        for block in blocks:
            self._add_block_to_document(doc, block, image_size)
    
    def _add_block_to_document(self, doc: Document, block: Dict, image_size: Tuple[int, int]) -> None:
        """Add a text block to the document."""
        text = block.get("text", "").strip()
        if not text:
            return
        
        bbox = block.get("bbox", [0, 0, 0, 0])
        confidence = block.get("confidence", 0)
        
        # Skip low confidence blocks
        if confidence < 30:
            logger.warning(f"Skipping low confidence block: {confidence}%")
            return
        
        # Determine block type based on characteristics
        block_type = self._classify_block(block, image_size)
        
        # Create paragraph with appropriate style
        para = doc.add_paragraph()
        
        # Set alignment
        alignment = self._determine_block_alignment(bbox, image_size)
        para.alignment = alignment
        
        # Add text with formatting based on block type
        run = para.add_run(text)
        self._apply_formatting(run, block_type, block)
        
        # Set paragraph spacing
        pf = para.paragraph_format
        if block_type == "title":
            pf.space_before = Pt(24)
            pf.space_after = Pt(12)
        elif block_type == "heading":
            pf.space_before = Pt(18)
            pf.space_after = Pt(6)
        else:
            pf.space_before = Pt(0)
            pf.space_after = Pt(6)
    
    def _classify_block(self, block: Dict, image_size: Tuple[int, int]) -> str:
        """Classify a text block as title, heading, or paragraph."""
        text = block.get("text", "")
        bbox = block.get("bbox", [0, 0, 0, 0])
        font_size = block.get("font_size", 12)
        
        # Height-based font size estimation
        height = bbox[3] - bbox[1]
        estimated_size = height * 0.75  # Approximate
        
        # Check for title characteristics
        if estimated_size > 24 or font_size > 20:
            return "title"
        
        # Check for heading characteristics
        if estimated_size > 16 or font_size > 14:
            return "heading"
        
        # Check for short, emphasized text
        word_count = len(text.split())
        if word_count <= 10 and estimated_size > 12:
            return "heading"
        
        return "paragraph"
    
    def _determine_block_alignment(self, bbox: List[float], image_size: Tuple[int, int]) -> int:
        """Determine text alignment from block position."""
        img_width = image_size[0]
        block_center = (bbox[0] + bbox[2]) / 2
        page_center = img_width / 2
        
        left_margin = bbox[0]
        right_margin = img_width - bbox[2]
        
        # Centered if block center is near page center
        if abs(block_center - page_center) < img_width * 0.1:
            return WD_ALIGN_PARAGRAPH.CENTER
        
        # Right aligned if right margin is small
        if right_margin < img_width * 0.1 and left_margin > img_width * 0.3:
            return WD_ALIGN_PARAGRAPH.RIGHT
        
        return WD_ALIGN_PARAGRAPH.LEFT
    
    def _apply_formatting(self, run, block_type: str, block: Dict) -> None:
        """Apply formatting to a run based on block type."""
        font = run.font
        
        if block_type == "title":
            font.size = Pt(24)
            font.bold = True
            font.name = "Arial"
        elif block_type == "heading":
            font.size = Pt(16)
            font.bold = True
            font.name = "Arial"
        else:
            font.size = Pt(11)
            font.bold = False
            font.name = "Arial"
        
        # Apply any detected formatting from OCR
        if block.get("is_bold"):
            font.bold = True
        if block.get("is_italic"):
            font.italic = True

