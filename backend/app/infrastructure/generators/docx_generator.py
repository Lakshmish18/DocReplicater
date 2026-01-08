"""
DOCX Generator
Generates DOCX documents from design schema and content sections.
"""

from typing import List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.domain.entities.design_schema import DesignSchema, StyleToken, FontWeight, TextAlignment
from app.domain.entities.content_section import ContentSection, SectionType
from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocxGenerator:
    """
    Generates DOCX documents with preserved formatting.
    
    Uses design schema to apply consistent styles and
    content sections to populate the document.
    """
    
    ALIGNMENT_MAP = {
        TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    
    def __init__(self, design_schema: DesignSchema):
        self.schema = design_schema
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self) -> None:
        """Set up document with page setup from schema."""
        section = self.doc.sections[0]
        page_setup = self.schema.page_setup
        
        # Set page dimensions
        section.page_width = Inches(page_setup.width)
        section.page_height = Inches(page_setup.height)
        
        # Set margins
        section.top_margin = Inches(page_setup.margin_top)
        section.bottom_margin = Inches(page_setup.margin_bottom)
        section.left_margin = Inches(page_setup.margin_left)
        section.right_margin = Inches(page_setup.margin_right)
        
        # Set up default styles
        self._setup_styles()
    
    def _setup_styles(self) -> None:
        """Set up document styles from schema style tokens."""
        for token_name, token in self.schema.style_tokens.items():
            try:
                self._create_or_update_style(token)
            except Exception as e:
                logger.warning(f"Failed to set up style {token_name}: {e}")
    
    def _create_or_update_style(self, token: StyleToken) -> None:
        """Create or update a paragraph style from token."""
        styles = self.doc.styles
        
        # Map token names to built-in styles
        style_mapping = {
            "Title": "Title",
            "H1": "Heading 1",
            "H2": "Heading 2",
            "H3": "Heading 3",
            "Body": "Normal",
            "Caption": "Caption",
            "Quote": "Quote",
        }
        
        style_name = style_mapping.get(token.name, token.name)
        
        try:
            style = styles[style_name]
        except KeyError:
            # Style doesn't exist, use Normal
            style = styles["Normal"]
        
        # Apply font properties
        font = style.font
        font.name = token.font.family
        font.size = Pt(token.font.size)
        font.bold = token.font.weight in [FontWeight.BOLD, FontWeight.SEMIBOLD, FontWeight.EXTRABOLD]
        font.italic = token.font.italic
        font.underline = token.font.underline
        
        # Apply color
        if token.font.color and token.font.color != "#000000":
            font.color.rgb = self._hex_to_rgb(token.font.color)
        
        # Apply paragraph properties
        pf = style.paragraph_format
        
        if token.alignment in self.ALIGNMENT_MAP:
            pf.alignment = self.ALIGNMENT_MAP[token.alignment]
        
        if token.space_before > 0:
            pf.space_before = Pt(token.space_before)
        if token.space_after > 0:
            pf.space_after = Pt(token.space_after)
        if token.first_line_indent != 0:
            pf.first_line_indent = Inches(token.first_line_indent)
        if token.left_indent != 0:
            pf.left_indent = Inches(token.left_indent)
        if token.right_indent != 0:
            pf.right_indent = Inches(token.right_indent)
    
    def generate(self, sections: List[ContentSection], output_path: str) -> str:
        """
        Generate DOCX from content sections.
        
        Args:
            sections: List of content sections in order
            output_path: Path to save the document
            
        Returns:
            Path to generated document
        """
        logger.info(f"Generating DOCX with {len(sections)} sections")
        
        for section in sections:
            self._add_section(section)
        
        # Ensure directory exists
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        self.doc.save(output_path)
        logger.info(f"DOCX saved to {output_path}")
        
        return output_path
    
    def _add_section(self, section: ContentSection) -> None:
        """Add a content section to the document."""
        if section.section_type == SectionType.TABLE:
            self._add_table(section)
        elif section.section_type in [SectionType.BULLET_LIST, SectionType.NUMBERED_LIST]:
            self._add_list(section)
        elif section.section_type == SectionType.PAGE_BREAK:
            self.doc.add_page_break()
        elif section.section_type == SectionType.IMAGE:
            self._add_image(section)
        else:
            self._add_paragraph(section)
    
    def _add_paragraph(self, section: ContentSection) -> None:
        """Add a paragraph to the document."""
        # Get style token
        token = self.schema.get_style_token(section.style_token)
        
        # Map section type to style
        style_mapping = {
            SectionType.TITLE: "Title",
            SectionType.HEADING_1: "Heading 1",
            SectionType.HEADING_2: "Heading 2",
            SectionType.HEADING_3: "Heading 3",
            SectionType.PARAGRAPH: "Normal",
            SectionType.CAPTION: "Caption",
            SectionType.QUOTE: "Quote",
        }
        
        style_name = style_mapping.get(section.section_type, "Normal")
        
        para = self.doc.add_paragraph(style=style_name)
        run = para.add_run(section.content)
        
        # Apply additional formatting from token if available
        if token:
            run.font.name = token.font.family
            run.font.size = Pt(token.font.size)
            run.font.bold = token.font.weight in [FontWeight.BOLD, FontWeight.SEMIBOLD, FontWeight.EXTRABOLD]
            run.font.italic = token.font.italic
            
            if token.font.color and token.font.color != "#000000":
                run.font.color.rgb = self._hex_to_rgb(token.font.color)
            
            # Apply paragraph formatting
            if token.alignment in self.ALIGNMENT_MAP:
                para.alignment = self.ALIGNMENT_MAP[token.alignment]
    
    def _add_table(self, section: ContentSection) -> None:
        """Add a table to the document."""
        if not section.table_data:
            return
        
        rows = len(section.table_data)
        cols = len(section.table_data[0]) if section.table_data else 0
        
        if rows == 0 or cols == 0:
            return
        
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        
        for row_idx, row_data in enumerate(section.table_data):
            for col_idx, cell_data in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = str(cell_data) if cell_data else ""
                
                # Bold headers
                if row_idx == 0 and section.table_headers:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _add_list(self, section: ContentSection) -> None:
        """Add a bulleted or numbered list."""
        is_numbered = section.section_type == SectionType.NUMBERED_LIST
        
        for idx, item in enumerate(section.list_items):
            para = self.doc.add_paragraph(style="List Bullet" if not is_numbered else "List Number")
            para.add_run(item)
    
    def _add_image(self, section: ContentSection) -> None:
        """Add an image to the document."""
        if not section.image_path:
            return
        
        try:
            # Add image with reasonable width
            self.doc.add_picture(section.image_path, width=Inches(5))
            
            # Add caption if available
            if section.image_alt_text:
                para = self.doc.add_paragraph(style="Caption")
                para.add_run(section.image_alt_text)
        except Exception as e:
            logger.warning(f"Failed to add image: {e}")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[Image: {section.image_alt_text or 'Image'}]")
    
    @staticmethod
    def _hex_to_rgb(hex_color: str) -> RGBColor:
        """Convert hex color to RGBColor."""
        hex_color = hex_color.lstrip('#')
        return RGBColor(
            int(hex_color[0:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16)
        )

