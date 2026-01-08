"""
DOCX Parser
Extracts design schema and content sections from DOCX files.
"""

from typing import List, Tuple, Optional
from pathlib import Path
from uuid import UUID
import mammoth
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from app.domain.entities.document import Document as DocumentEntity
from app.domain.entities.design_schema import (
    DesignSchema, StyleToken, PageSetup, FontStyle, FontWeight, TextAlignment
)
from app.domain.entities.content_section import ContentSection, SectionType
from app.utils.logger import get_logger

logger = get_logger(__name__)


class DocxParser:
    """
    Parser for DOCX files.
    
    Extracts:
    - Complete design schema (styles, formatting, page setup)
    - Content sections with style references
    - Tables and lists
    """
    
    # Map Word alignment to our alignment enum
    ALIGNMENT_MAP = {
        WD_ALIGN_PARAGRAPH.LEFT: TextAlignment.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER: TextAlignment.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT: TextAlignment.RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY: TextAlignment.JUSTIFY,
    }
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = Document(file_path)
        self._style_cache = {}
    
    def parse(self, document_id: UUID) -> Tuple[DesignSchema, List[ContentSection]]:
        """
        Parse the DOCX file and extract design schema and content sections.
        
        Returns:
            Tuple of (DesignSchema, List[ContentSection])
        """
        logger.info(f"Parsing DOCX: {self.file_path}")
        
        # Extract design schema
        design_schema = self._extract_design_schema(document_id)
        
        # Extract content sections
        sections = self._extract_content_sections(document_id)
        
        logger.info(f"Extracted {len(sections)} sections from DOCX")
        
        return design_schema, sections
    
    def _extract_design_schema(self, document_id: UUID) -> DesignSchema:
        """Extract the complete design schema from the document."""
        
        # Extract page setup
        page_setup = self._extract_page_setup()
        
        # Extract style tokens
        style_tokens = self._extract_style_tokens()
        
        # Determine heading hierarchy
        heading_hierarchy = self._infer_heading_hierarchy()
        
        # Extract color palette
        color_palette = self._extract_color_palette()
        
        schema = DesignSchema(
            document_id=document_id,
            page_setup=page_setup,
            style_tokens=style_tokens,
            default_font_family=self._get_default_font(),
            default_font_size=self._get_default_font_size(),
            heading_hierarchy=heading_hierarchy,
            color_palette=color_palette,
            extracted_from="docx",
            confidence_score=1.0,
        )
        
        # Lock the schema to prevent modifications
        schema.lock()
        
        return schema
    
    def _extract_page_setup(self) -> PageSetup:
        """Extract page setup from document sections."""
        section = self.doc.sections[0]
        
        return PageSetup(
            width=section.page_width.inches if section.page_width else 8.5,
            height=section.page_height.inches if section.page_height else 11.0,
            orientation="portrait" if section.page_width < section.page_height else "landscape",
            margin_top=section.top_margin.inches if section.top_margin else 1.0,
            margin_bottom=section.bottom_margin.inches if section.bottom_margin else 1.0,
            margin_left=section.left_margin.inches if section.left_margin else 1.0,
            margin_right=section.right_margin.inches if section.right_margin else 1.0,
        )
    
    def _extract_style_tokens(self) -> dict:
        """Extract all paragraph styles as tokens."""
        tokens = {}
        
        for style in self.doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                try:
                    token = self._style_to_token(style)
                    if token:
                        tokens[token.name] = token
                        self._style_cache[style.name] = token.name
                except Exception as e:
                    logger.warning(f"Failed to extract style {style.name}: {e}")
        
        # Ensure we have default tokens
        default_tokens = DesignSchema.create_default_tokens()
        for name, token in default_tokens.items():
            if name not in tokens:
                tokens[name] = token
        
        return tokens
    
    def _style_to_token(self, style) -> Optional[StyleToken]:
        """Convert a Word style to a StyleToken."""
        if not style.name:
            return None
        
        # Extract font properties
        font = style.font
        font_style = FontStyle(
            family=font.name or "Arial",
            size=font.size.pt if font.size else 12.0,
            weight=FontWeight.BOLD if font.bold else FontWeight.NORMAL,
            italic=font.italic or False,
            underline=font.underline or False,
            color=self._rgb_to_hex(font.color.rgb) if font.color and font.color.rgb else "#000000",
        )
        
        # Extract paragraph properties
        pf = style.paragraph_format
        alignment = TextAlignment.LEFT
        if pf.alignment and pf.alignment in self.ALIGNMENT_MAP:
            alignment = self.ALIGNMENT_MAP[pf.alignment]
        
        return StyleToken(
            name=self._normalize_style_name(style.name),
            font=font_style,
            alignment=alignment,
            line_spacing=pf.line_spacing if pf.line_spacing else 1.15,
            space_before=pf.space_before.pt if pf.space_before else 0.0,
            space_after=pf.space_after.pt if pf.space_after else 0.0,
            first_line_indent=pf.first_line_indent.inches if pf.first_line_indent else 0.0,
            left_indent=pf.left_indent.inches if pf.left_indent else 0.0,
            right_indent=pf.right_indent.inches if pf.right_indent else 0.0,
        )
    
    def _normalize_style_name(self, name: str) -> str:
        """Normalize style names to standard tokens."""
        name_lower = name.lower()
        
        if "title" in name_lower:
            return "Title"
        elif "heading 1" in name_lower or name_lower == "h1":
            return "H1"
        elif "heading 2" in name_lower or name_lower == "h2":
            return "H2"
        elif "heading 3" in name_lower or name_lower == "h3":
            return "H3"
        elif "heading" in name_lower:
            return "H2"
        elif "body" in name_lower or "normal" in name_lower:
            return "Body"
        elif "caption" in name_lower:
            return "Caption"
        elif "quote" in name_lower:
            return "Quote"
        
        return name
    
    def _extract_content_sections(self, document_id: UUID) -> List[ContentSection]:
        """Extract content sections from the document."""
        sections = []
        order_index = 0
        
        for para in self.doc.paragraphs:
            if not para.text.strip():
                continue
            
            section_type = self._determine_section_type(para)
            style_token = self._get_style_token(para)
            
            section = ContentSection(
                document_id=document_id,
                order_index=order_index,
                section_type=section_type,
                content=para.text,
                original_content=para.text,
                style_token=style_token,
                editable=True,
                ai_enabled=True,
            )
            
            sections.append(section)
            order_index += 1
        
        # Extract tables
        for table_idx, table in enumerate(self.doc.tables):
            table_data = []
            headers = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text for cell in row.cells]
                if row_idx == 0:
                    headers = row_data
                table_data.append(row_data)
            
            section = ContentSection(
                document_id=document_id,
                order_index=order_index,
                section_type=SectionType.TABLE,
                content="",
                original_content="",
                style_token="Body",
                table_data=table_data,
                table_headers=headers,
                editable=True,
                ai_enabled=False,
            )
            
            sections.append(section)
            order_index += 1
        
        return sections
    
    def _determine_section_type(self, para) -> SectionType:
        """Determine the section type from paragraph style."""
        style_name = para.style.name.lower() if para.style else ""
        
        if "title" in style_name:
            return SectionType.TITLE
        elif "heading 1" in style_name:
            return SectionType.HEADING_1
        elif "heading 2" in style_name:
            return SectionType.HEADING_2
        elif "heading 3" in style_name:
            return SectionType.HEADING_3
        elif "list" in style_name:
            if "bullet" in style_name or "number" not in style_name:
                return SectionType.BULLET_LIST
            return SectionType.NUMBERED_LIST
        elif "quote" in style_name:
            return SectionType.QUOTE
        else:
            return SectionType.PARAGRAPH
    
    def _get_style_token(self, para) -> str:
        """Get the style token name for a paragraph."""
        if para.style and para.style.name in self._style_cache:
            return self._style_cache[para.style.name]
        
        # Fallback based on section type
        section_type = self._determine_section_type(para)
        type_to_token = {
            SectionType.TITLE: "Title",
            SectionType.HEADING_1: "H1",
            SectionType.HEADING_2: "H2",
            SectionType.HEADING_3: "H3",
            SectionType.PARAGRAPH: "Body",
            SectionType.QUOTE: "Quote",
            SectionType.CAPTION: "Caption",
        }
        return type_to_token.get(section_type, "Body")
    
    def _infer_heading_hierarchy(self) -> List[str]:
        """Infer heading hierarchy from document."""
        hierarchy = []
        
        for style in self.doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                name_lower = style.name.lower()
                if "heading" in name_lower:
                    hierarchy.append(self._normalize_style_name(style.name))
        
        # Sort by heading level
        def heading_sort_key(h):
            if h == "Title":
                return 0
            try:
                return int(h.replace("H", ""))
            except:
                return 99
        
        return sorted(set(hierarchy), key=heading_sort_key)
    
    def _extract_color_palette(self) -> List[str]:
        """Extract unique colors used in the document."""
        colors = set()
        
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.color and run.font.color.rgb:
                    colors.add(self._rgb_to_hex(run.font.color.rgb))
        
        return list(colors)
    
    def _get_default_font(self) -> str:
        """Get the default font family."""
        try:
            default_style = self.doc.styles["Normal"]
            if default_style.font.name:
                return default_style.font.name
        except:
            pass
        return "Arial"
    
    def _get_default_font_size(self) -> float:
        """Get the default font size."""
        try:
            default_style = self.doc.styles["Normal"]
            if default_style.font.size:
                return default_style.font.size.pt
        except:
            pass
        return 12.0
    
    @staticmethod
    def _rgb_to_hex(rgb) -> str:
        """Convert RGB to hex color."""
        if rgb is None:
            return "#000000"
        return "#{:02x}{:02x}{:02x}".format(rgb[0], rgb[1], rgb[2])
    
    def get_html_preview(self) -> str:
        """Generate HTML preview using mammoth."""
        try:
            with open(self.file_path, "rb") as f:
                result = mammoth.convert_to_html(f)
                return result.value
        except Exception as e:
            logger.error(f"Failed to generate HTML preview: {e}")
            return ""

